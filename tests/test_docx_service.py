import pytest
from docx import Document
from services.docx_service import DOCXService

def test_load_template():
    service = DOCXService()
    # This test requires a template to exist
    # Skip if template not available
    try:
        doc = service.load_template("PENGADAAN", "RKS")
        assert isinstance(doc, Document)
    except FileNotFoundError:
        pytest.skip("Template not found")

def test_replace_placeholder():
    service = DOCXService()
    doc = Document()
    doc.add_paragraph("Project: {{project_name}}")

    replacements = {"project_name": "Test Project"}
    service.replace_placeholders(doc, replacements)

    text = doc.paragraphs[0].text
    assert "Test Project" in text
    assert "{{project_name}}" not in text
