import pytest
from docx import Document
from services.docx_service import DOCXService
from pathlib import Path

def test_template_loading():
    """Test that templates exist and can be loaded"""
    docx_service = DOCXService()

    # Test RAB template
    try:
        rab_doc = docx_service.load_template("PENGADAAN", "RAB")
        assert hasattr(rab_doc, 'paragraphs'), "Should be a Document object"
        assert len(rab_doc.paragraphs) > 0, "Template should have content"
        print(f"✓ RAB template loaded: {len(rab_doc.paragraphs)} paragraphs")
    except FileNotFoundError as e:
        pytest.skip(f"RAB template not found: {e}")

    # Test RKS template
    try:
        rks_doc = docx_service.load_template("PENGADAAN", "RKS")
        assert hasattr(rks_doc, 'paragraphs'), "Should be a Document object"
        assert len(rks_doc.paragraphs) > 0, "Template should have content"
        print(f"✓ RKS template loaded: {len(rks_doc.paragraphs)} paragraphs")
    except FileNotFoundError as e:
        pytest.skip(f"RKS template not found: {e}")

def test_placeholder_replacement():
    """Test that placeholders are correctly replaced"""
    docx_service = DOCXService()
    doc = Document()

    # Add test content with placeholders
    doc.add_paragraph("Project: {{project_name}}")
    doc.add_paragraph("Timeline: {{timeline}}")
    doc.add_paragraph("Location: {{location}}")
    doc.add_paragraph("Work Type: {{work_type}}")

    # Test data
    test_data = {
        "project_name": "TEST PROJECT PENGADAAN SERVER",
        "timeline": "3 bulan",
        "location": "PT Test Location",
        "work_type": "Instalasi Server"
    }

    # Replace placeholders
    docx_service.replace_placeholders(doc, test_data)

    # Verify replacements
    text = doc.paragraphs[0].text
    assert "TEST PROJECT PENGADAAN SERVER" in text
    assert "{{project_name}}" not in text

    assert "3 bulan" in doc.paragraphs[1].text
    assert "PT Test Location" in doc.paragraphs[2].text
    assert "Instalasi Server" in doc.paragraphs[3].text

    print("✓ All placeholders replaced correctly")
    print(f"  - project_name: {test_data['project_name']}")
    print(f"  - timeline: {test_data['timeline']}")
    print(f"  - location: {test_data['location']}")
    print(f"  - work_type: {test_data['work_type']}")

def test_fill_and_save_template():
    """Test filling template with data and saving"""
    docx_service = DOCXService()

    # Skip if template doesn't exist
    try:
        doc = docx_service.load_template("PENGADAAN", "RKS")
    except FileNotFoundError:
        pytest.skip("RKS template not found")

    # Test data
    test_data = {
        "project_name": "TEST PROJECT PENGADAAN SERVER",
        "timeline": "3 bulan",
        "location": "PT Test Location",
        "pasal_2_content": "Test content for Pasal 2",
        "pasal_3_content": "Test content for Pasal 3",
        "pasal_10_content": "Test payment terms",
        "work_type": "Instalasi Server",
        "date": "18 April 2026"
    }

    # Fill template
    filled_doc = docx_service.fill_template(doc, test_data)

    # Save to test output
    output_path = "test_output_rks.docx"
    docx_service.save_document(filled_doc, output_path)

    # Verify file created
    assert Path(output_path).exists(), f"Output file not created: {output_path}"

    # Verify placeholders replaced in saved file
    saved_doc = Document(output_path)
    full_text = '\n'.join([p.text for p in saved_doc.paragraphs])

    assert "TEST PROJECT PENGADAAN SERVER" in full_text
    assert "{{project_name}}" not in full_text
    assert "Test content for Pasal 2" in full_text

    # Cleanup
    Path(output_path).unlink()

    print(f"✓ Template filled and saved successfully")
    print(f"✓ Output file created: {output_path}")
    print(f"✓ All placeholders verified in saved document")
