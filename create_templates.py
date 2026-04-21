from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def set_page_margins(section):
    """Set page margins in cm"""
    section.top_margin = Inches(3 / 2.54)    # 3cm
    section.bottom_margin = Inches(3 / 2.54)  # 2cm
    section.left_margin = Inches(1.5 / 2.54)   # 1.5cm
    section.right_margin = Inches(1 / 2.54)    # 1cm

def set_font(paragraph, font_name='Times New Roman', size=Pt(12)):
    """Set font for all runs in paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = size

def add_pasal_heading(doc, pasal_num, title):
    """Add a centered, bold, capitalized pasal heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'PASAL {pasal_num}\n{title.upper()}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def add_underline(run):
    """Add underline to run"""
    from docx.oxml.ns import qn
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.underline = True

def add_numbered_paragraph(doc, text):
    """Add a numbered paragraph"""
    p = doc.add_paragraph(text, style='List Number')
    set_font(p)
    return p

def create_rab_template():
    doc = Document()
    set_page_margins(doc.sections[0])

    # Main Title - centered, bold, capital, underline
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RENCANA ANGGARAN BIAYA')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.all_caps = True
    run.font.underline = True

    # Subtitle - nama pekerjaan centered under title
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{project_name}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

    # Project Info
    doc.add_paragraph()
    p = doc.add_paragraph('Waktu Pelaksanaan: {{timeline}}')
    set_font(p)

    # Items table
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Uraian'
    hdr_cells[2].text = 'Volume'
    hdr_cells[3].text = 'Satuan'
    hdr_cells[4].text = 'Harga Satuan (IDR)'
    hdr_cells[5].text = 'Jumlah Harga (IDR)'

    # Center header text and make bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

    # Sample data row
    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = '{{items}}'
    row_cells[2].text = ''
    row_cells[3].text = ''
    row_cells[4].text = ''
    row_cells[5].text = ''

    # Total rows inside table
    # Jumlah row
    total_row_cells = table.add_row().cells
    total_row_cells[0].merge(total_row_cells[4]).text = 'Jumlah'
    total_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row_cells[5].text = 'Rp {{subtotal}}'

    # PPN row
    ppn_row_cells = table.add_row().cells
    ppn_row_cells[0].merge(ppn_row_cells[4]).text = 'PPN 11%'
    ppn_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ppn_row_cells[5].text = 'Rp {{ppn}}'

    # Jumlah Total row
    final_row_cells = table.add_row().cells
    final_row_cells[0].merge(final_row_cells[4]).text = 'Jumlah Total'
    final_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in final_row_cells[0].paragraphs[0].runs:
        run.font.bold = True
    final_row_cells[5].text = 'Rp {{total}}'
    for run in final_row_cells[5].paragraphs[0].runs:
        run.font.bold = True

    # Location and date at end - aligned right
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('Dikeluarkan di: {{location}}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p)

    p = doc.add_paragraph('Tanggal: {{date}}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p)

    Path("templates").mkdir(exist_ok=True)
    doc.save("templates/RAB_pengadaan.docx")
    print("Created: templates/RAB_pengadaan.docx")

def create_rks_template():
    doc = Document()
    set_page_margins(doc.sections[0])

    # Main Title - centered, bold, capital, underline
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RENCANA KERJA DAN SYARAT')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.all_caps = True
    run.font.underline = True

    # Subtitle - nama pekerjaan centered under title
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{project_name}}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

    # Project Info
    doc.add_paragraph()
    p = doc.add_paragraph('Waktu Pelaksanaan: {{timeline}}')
    set_font(p)

    # Static pasal (no extraction needed)
    doc.add_paragraph()
    add_pasal_heading(doc, 1, 'MAKSUD DAN TUJUAN')
    p = doc.add_paragraph('Syarat-syarat teknis atau spesifikasi teknis ini dimaksudkan sebagai pedoman teknis dan acuan dalam pelaksanaan pekerjaan {{project_name}}.\n\nAdapun tujuan dari syarat-syarat teknis atau spesifikasi teknis ini adalah menghasilkan standar kualitas perangkat dan performa layanan yang memadai guna pendukung kebutuhan operasional.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 2, 'POKOK-POKOK DAN JENIS PEKERJAAN')
    p = doc.add_paragraph('{{pasal_2_content}}')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 3, 'RUANG LINGKUP PEKERJAAN')
    p = doc.add_paragraph('{{pasal_3_content}}')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 4, 'JENIS KONTRAK')
    p = doc.add_paragraph('1. Skema Kontrak adalah berupa harga barang dengan jenis Kontrak Lumsum.\n2. Harga yang dimaksud pada butir (1) Pasal ini adalah merupakan harga yang tetap (Fixed Price) selama masa berlakunya kontrak.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 5, 'GARANSI')
    p = doc.add_paragraph('1. Kontraktor wajib memberikan Garansi atas setiap pekerjaan yang dilakukannya dan atau untuk perkerjaan perbaikan, perubahan, dan atau peningkatan kualitas Produk.\n2. Garansi sebagaimana termaktub dalam butir 1 pasal ini, mencakup:\na. Kondisi fisik, fungsionalitas dan operasionalitas Produk;\nb. Pelaksanaan pekerjaan yang wajar dan profesional setiap saat;\nc. Memberikan pelayanan dengan keahlian standar, berhati – hati dan tekun;\nd. Menyiapkan personil yang memadai, dengan keahlian yang cukup, untuk memberikan pelayanan terbaik kepada Pemberi Tugas.\n3. Waktu pelaksanaan garansi pekerjaan sesuai dengan butir 1 pasal ini adalah ditetapkan selama 12 (dua belas) bulan sejak Berita Acara Penyelesaian Pekerjaan Tahap 1 ditanda tangani oleh kedua belah pihak.\n4. Masa garansi produk/barang berlaku 12 (dua belas) bulan sejak Berita Acara Serah Terima atau Surat Jalan diterima dan ditanda tangani oleh Pihak PT TPS.\n5. Garansi ini tidak berlaku untuk kerusakan Produk dan alat pendukungnya karena wear and tear atau secara elektronik karena penggunaan.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 6, 'WAKTU PELAKSANAAN PEKERJAAN')
    p = doc.add_paragraph('Waktu pelaksanaan pekerjaan {{project_name}} adalah ditetapkan selama {{timeline}} sejak PO terbit.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 7, 'TEMPAT PELAKSANAAN DAN JADWAL PEKERJAAN')
    p = doc.add_paragraph('Melakukan pengiriman barang di jam kerja (Senin s.d. Jumat) mulai jam 08.00 s/d. 16.00 WIB diruang kerja Departemen TI {{location}}.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 8, 'TEMPAT PENYERAHAN PEKERJAAN')
    p = doc.add_paragraph('1. Tempat penyerahan hasil pekerjaan adalah di {{location}}.\n2. Penyerahan hasil pekerjaan dilakukan setelah pekerjaan dinyatakan selesai dan memenuhi ketentuan sesuai dengan spesifikasi teknis.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 9, 'ADMINISTRASI DAN PELAPORAN')
    p = doc.add_paragraph('1. Kontraktor harus membuat laporan dan dokumentasi berdasarkan kegiatan yang dilakukan.\n2. Format-format dokumen yang akan dipergunakan dalam pelaksanaan pekerjaan harus mendapat persetujuan terlebih dahulu dari Pemberi Tugas.\n3. Laporan sebagaimana dimaksud diatas, dapat disampaikan dengan cara presentasi jika diperlukan.\n4. Semua laporan tersebut diatas harus dilengkapi dengan Berita Acara Penyelesaian Pekerjaan serta invoice dan faktur pajak untuk dilampirkan dalam surat permohonan pembayaran.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 10, 'CARA PEMBAYARAN')
    p = doc.add_paragraph('{{pasal_10_content}}')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 11, 'KEWAJIBAN DAN HAK PEMBERI TUGAS DAN PEMASOK BARANG')
    p = doc.add_paragraph('Dalam pelaksanaan pekerjaan pekerjaan {{project_name}}, masing-masing pihak mempunyai kewajiban dan hak yang diatur sebagai berikut:\n\n1. Kewajiban Pemberi Tugas:\na. Melakukan pembayaran kepada Kontraktor sesuai kontrak (surat perjanjian).\nb. Menyediakan Tanda Pengenal bagi seluruh tenaga kerja yang dipekerjakan oleh kontraktor.\n\n2. Kewajiban Kontraktor:\na. Kontraktor pelaksana pekerjaan diwajibkan memenuhi dan melakukan pekerjaan seperti yang tertuang pada pasal 3 (tiga) tentang ruang lingkup pekerjaan;\nb. Kontraktor pelaksana pekerjaan wajib menyelesaikan beban pekerjaan seperti yang tertuang pada pasal 3 (tiga) tentang Ruang Lingkup Pekerjaan selama kurun waktu yang telah diatur pada pasal 6 (enam) tentang Waktu Pelaksanaan Pekerjaan;\nc. Kontraktor pelaksana pekerjaan dan tenaga kerja yang ditugaskan diwajibkan memiliki pengalaman dalam project {{work_type}} dilingkungan {{location}};\nd. Kontraktor pelaksana pekerjaan diwajibkan memiliki surat dukungan dari principal atas produk yang tercantum pada Tabel 3.1 tentang Lingkup Item pekerjaan {{project_name}};\ne. Kontraktor pelaksana pekerjaan dimaksud diutamakan memiliki alamat domisili atau kantor cabang di area Surabaya Raya yang ditunjukkan dengan Surat Keterangan Domisili yang berlaku guna kemudahan pelaksanaan koordinasi dan komunikasi serta dukungan layanan (warranty/replacement);\nf. Menyediakan semua kebutuhan peralatan (tools) dan perlengkapan kerja serta peralatan bantu yang diperlukan dalam jumlah yang cukup;\ng. Melakukan pemeliharaan dan perbaikan terhadap prasarana, sarana dan peralatan yang disediakan oleh Pemberi Tugas, termasuk pembersihan maupun perbaikan bila terjadi kerusakan;\nh. Menjaga kerapian, keindahan, ketertiban dan kebersihan ruang kerja atau area kerja yang dipergunakan;\ni. Membantu menyiapkan telahan dan memberikan saran tertulis kepada Pemberi Tugas dalam rangka meningkatkan kinerja perangkat beserta peripheral pada lingkup pelaksanaan pekerjaan {{project_name}};\nj. Menyediakan peralatan keselamatan kerja (K3) sesuai ketentuan dan peraturan yang berlaku kepada setiap tenaga kerja yang ditempatkan, sebagai kelengkapan dalam pelaksanaan tugas;\nk. Menempatkan tenaga kerja yang terampil dan mengganti tenaga kerja yang tidak terampil, tidak disiplin serta melanggar peraturan yang berlaku di PT. TPS;\nl. Menjaga agar para tenaga kerja yang ditempatkan setiap bertugas harus menggunakan pakaian seragam dan peralatan keselamatan kerja (K3) sebagaimana mestinya yang disediakan oleh Kontraktor;\nm. Menjaga agar para tenaga kerja yang ditempatkan memakai tanda pengenal diri yang jelas yang diterbitkan oleh Pemberi Tugas;\nn. Menjaga agar para tenaga kerja yang ditempatkan memiliki tingkat disiplin kerja yang tinggi dan kemampuan komunikasi yang baik;\no. Menjaga agar para tenaga kerja yang ditempatkan harus senantiasa mentaati semua peraturan dan ketentuan yang berlaku di PT. TPS;\np. Bertanggung jawab sepenuhnya apabila terjadi tuntutan atau gugatan hukum yang diajukan oleh para tenaga kerja yang ditempatkan oleh Kontraktor maupun oleh Pihak lain yang merasa dirugikan oleh Kontraktor;\nq. Bertanggung jawab sepenuhnya atas pelanggaran hukum yang terjadi di PT TPS yang dilakukan oleh para tenaga kerja yang ditempatkan;\nr. Bertanggung jawab atas semua pelanggaran ketentuan larangan yang dilakukan oleh para tenaga kerja yang ditempatkan di PT TPS;\ns. Bertanggung jawab terhadap semua kerugian yang mungkin timbul akibat perbuatan tenaga kerja yang ditempatkan oleh Kontraktor, baik sengaja maupun tidak sengaja dan atau atas kelalaian Pihak lain, sehingga mengakibatkan rusaknya alat, fasilitas dan peralatan kerja dengan tanpa alasan apapun;\nt. Bertanggung jawab atas semua biaya perawatan dari para tenaga kerja, orang lain maupun Pihak lain yang menderita sakit atau cidera akibat perbuatan atau kelalaian dari tenaga kerja yang ditempatkan kontraktor selama menjalankan tugas pekerjaannya;\nu. Apabila terjadi kecelakaan kerja yang mengakibatkan meninggalnya tenaga kerja yang ditempatkan ataupun orang lain, Kontraktor bertanggung jawab penuh terhadap semua pengurusan jenasah dan biaya-biaya yang diperlukan kepada ahli warisnya.\n\n3. Hak-hak Pemberi Tugas:\na. Memerintah kontraktor dan atau kepada para tenaga kerja yang ditempatkan untuk melaksanakan pekerjaan {{project_name}} dengan sebaik-baiknya sesuai dengan ruang lingkup pekerjaan;\nb. Melakukan tegoran-tegoran atau larangan secara lisan/tertulis kepada Kontraktor dan atau kepada para tenaga kerja yang ditempatkan apabila di dalam pelaksanaan tugasnya ternyata terjadi penyimpangan-penyimpangan dan atau pelanggaran-pelanggaran lainnya;\nc. Meminta penggantian tenaga kerja yang ditempatkan oleh kontraktor, apabila kenyataan di lapangan menunjukkan tenaga kerja yang bersangkutan dinilai tidak terampil dan tidak dapat melaksanakan kewajiban dan tugasnya dengan baik;\nd. Meneliti kebenaran laporan-laporan yang disampaikan oleh kontraktor, yang diantaranya akan dipergunakan sebagai dasar permohonan pembayaran.\n\n4. Hak-hak Kontraktor:\nMenerima pembayaran dari Pemberi Tugas sesuai dengan hasil pekerjaan yang telah selesai dilaksanakan dengan hasil baik.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 12, 'KESELAMATAN, KESEHATAN, KEAMANAN DAN LINGKUNGAN KERJA')
    p = doc.add_paragraph('1. Tenaga Kerja Kontraktor harus telah mendapatkan Induction Training mengenai Keselamatan, Kesehatan, Keamanan dan Lingkungan Kerja sebelum bekerja di Area Terbatas PT. TPS.\n2. Kontraktor harus menyediakan kendaraan operasional untuk kegiatan yang berkaitan dengan pelaksanaan Perjanjian ini di Area Terbatas PT. TPS apabila diperlukan.\n3. Kontraktor berkewajiban menyediakan peralatan pelindung diri (APD) bagi karyawannya serta mentaati peraturan Keselamatan, Kesehatan, Keamanan dan Lingkungan Kerja (LK3M).\n4. Kontraktor harus selalu menjaga ketertiban, kerapian dan kebersihan lingkungan kerja dan mematuhi peraturan yang berlaku pada PT. TPS.\n5. Kontraktor harus mendukung PT. TPS dalam rangka implementasi dan pemenuhan kebutuhan audit ISO 9001, ISO 45001, ISO 28000, ISO 27001 serta kebutuhan audit-audit lainnya yang berlaku di PT TPS.\n6. Kontraktor bertanggung jawab terhadap keamanan, keselamatan dan segala resiko yang terjadi pada karyawannya yang dipekerjakan di PT. TPS.\n7. Apabila terjadi kecelakaan kerja di lingkungan PT. TPS, para tenaga kerja yang ditempatkan oleh Kontraktor wajib membantu dengan sukarela tanpa mengganggu pelaksanaan tugas yang menjadi tanggung jawabnya.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 13, 'KOMITMEN KERAHASIAAN')
    p = doc.add_paragraph('1. Setiap tenaga kerja yang ditugaskan wajib mengisi dan menandatangai form NDA (non-disclosure agreement) sebagai pernyataan tertulis terhadap komitmen dalam menjaga kerahasiaan data dan atau informasi milik PT. TPS.\n2. Kontraktor berkewajiban untuk melindungi Hak Cipta dari segala gagasan, ide, know-how, dan teknik – teknik yang digunakan dan/ atau tercipta akibat pelaksanaan pekerjaan.\n3. Kontraktor berkewajiban untuk menjaga kerahasiaan setiap informasi yang didapat dari Pengguna Jasa dan/ atau segala dokumen – dokumen kepada siapapun, dan tidak diperkenankan untuk membuatnya terbuka untuk umum dengan tujuan pengiklanan atau apapun itu, tanpa persetujuan tertulis dari Pengguna Jasa.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 14, 'PERATURAN DAN KETENTUAN YANG BERLAKU')
    p = doc.add_paragraph('Peraturan dan ketentuan yang harus dipatuhi oleh Kontraktor dalam melaksanakan pekerjaan diantaranya adalah:\n1. Peraturan-peraturan dan ketentuan dalam ketenaga-kerjaan atau hukum perburuhan yang berlaku di Indonesia.\n2. Peraturan-peraturan dan ketentuan yang berlaku di PT. TPS (misalnya peraturan keselamatan).\n3. Peraturan-peraturan dan ketentuan-ketentuan lain yang ada hubungannya dengan pekerjaan ini.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 15, 'KONDISI TIDAK NORMAL')
    p = doc.add_paragraph('Kontraktor berkewajiban menanggung biaya perbaikan dan penggantian Suku cadang (spare parts) yang timbul akibat kelalaian maupun accident (kecelakaan) yang dilakukan oleh personel atau tenaga kerja Kontraktor.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 16, 'PENALTI')
    p = doc.add_paragraph('1. Apabila terjadi keterlambatan dalam penyelesaian waktu pelaksanaan pekerjaan sesuai yang dipersyaratkan pada pasal 6 (enam), maka kontraktor akan dikenakan denda sebesar 1 ‰ (satu permil) perhari keterlambatan dari jumlah harga borongan, dengan besaran denda maksimal sebesar 5% dari harga pekerjaan yang dinyatakan dalam Berita Acara.\n2. Keterlambatan akibat adanya force majeure tidak dikenakan denda.')
    set_font(p)

    doc.add_paragraph()
    add_pasal_heading(doc, 17, 'LAIN-LAIN')
    p = doc.add_paragraph('1. Apabila terjadi kecelakaan kerja yang diakibatkan oleh faktor non teknis terhadap alat yang penyelesaiannya ditangani oleh PT. TPS, maka Pelaksana Pekerjaan harus membantu semaksimal mungkin agar supaya alat dapat beroperasi sesegera mungkin dan tidak mempengaruhi kesiapan, keandalan alat serta mengganggu operasional keseluruhan di lingkungan PT.TPS.\n2. Pelaksana pekerjaan harus menjalin kerjasama yang baik dengan semua pihak yang ada di lingkungan PT. TPS secara professional berdasarkan peraturan-peraturan yang berlaku di lingkungan PT. TPS.\n3. Segala sesuatu yang belum tercantum dan diatur dalam syarat-syarat teknis ini dan nantinya ternyata dalam pelaksanaannya ada kaitannya dengan kontrak pekerjaan ini akan diatur dan ditentukan kemudian.')
    set_font(p)

    # Location and date at end - aligned right
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph('Dikeluarkan di: {{location}}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p)

    p = doc.add_paragraph('Tanggal: {{date}}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p)

    Path("templates").mkdir(exist_ok=True)
    doc.save("templates/RKS_pengadaan.docx")
    print("Created: templates/RKS_pengadaan.docx")

if __name__ == "__main__":
    create_rab_template()
    create_rks_template()
    print("Templates created successfully!")

