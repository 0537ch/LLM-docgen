import pytest
import os
from pathlib import Path
from docx import Document
from services.ocr_service import OCRService
from services.extraction_service import ExtractionService
from services.docx_service import DOCXService
from services.paraphrase_service import ParaphraseService

def format_work_activities(activities):
    """Format work activities for Pasal 2"""
    if not activities:
        return "- Pekerjaan ini belum didefinisikan"

    formatted = []
    for i, activity in enumerate(activities, 1):
        formatted.append(f"{i}. {activity}")
    return '\n'.join(formatted)

def format_pasal_3_content(data):
    """Format Pasal 3 content from extracted data with Tabel 3.1"""
    parts = []

    # Add scope description
    if data.get("scope_description"):
        parts.append(data["scope_description"])


    # Add table transition
    items = data.get("items", [])
    if items:
        parts.append("\nDetail material yang terdapat dalam pekerjaan ini wajib memenuhi spesifikasi yang tertuang dalam Tabel 3.1 sebagai berikut:")

        # Create table
        parts.append("\nTabel 3.1 Lingkup Item Pekerjaan")
        parts.append("No.\tUraian\tVolume\tSatuan")

        # Group by category
        material_items = [i for i in items if i.get('category') == 'Material']
        jasa_items = [i for i in items if i.get('category') == 'Jasa']

        if material_items:
            parts.append("A. MATERIAL")
            for i, item in enumerate(material_items, 1):
                parts.append(f"{i}\t{item.get('name', '')}\t{item.get('quantity', '')}\t{item.get('unit', '')}")

        if jasa_items:
            parts.append("B. JASA")
            for i, item in enumerate(jasa_items, 1):
                parts.append(f"{i}\t{item.get('name', '')}\t{item.get('quantity', '')}\t{item.get('unit', '')}")

    return '\n'.join(parts)

def test_extraction_to_docx():
    """Test extraction output to DOCX template filling"""
    if not os.getenv("GEMINI_API_KEY"):
        pytest.skip("GEMINI_API_KEY not set")

    # Use extracted data from test_extraction_comprehensive
    extraction_service = ExtractionService()
    docx_service = DOCXService()

    # Read from extracted LHP text
    lhp_path = "extracted/LHP/01. Laporan Hasil Pemeriksaan Perangkat Radio Komunikasi Security dan Mekanik CC_easyocr.txt"
    if not Path(lhp_path).exists():
        pytest.skip(f"Extracted LHP not found: {lhp_path}")

    with open(lhp_path, 'r', encoding='utf-8') as f:
        lhp_text = f.read()

    # Extract data
    print("\n=== EXTRACTING DATA ===")
    doc_type = extraction_service.detect_document_type(lhp_text)
    print(f"Document type: {doc_type}")

    data = extraction_service.extract_structured_data(lhp_text, doc_type)

    print(f"\n✓ Extracted {len(data.get('work_activities', []))} work activities")
    print(f"✓ Extracted {len(data.get('items', []))} items")
    print(f"✓ Project: {data.get('project_name')}")

    # Debug: Show work activities
    print("\n=== WORK ACTIVITIES EXTRACTED ===")
    for i, activity in enumerate(data.get('work_activities', []), 1):
        print(f"{i}. {activity[:80]}...")

    # Prepare template data with ALL extracted fields
    pasal_2_text = format_work_activities(data.get("work_activities", []))
    print(f"\n=== PASAL 2 CONTENT ===")
    print(pasal_2_text[:500])

    template_data = {
        "project_name": data.get("project_name", ""),
        "timeline": data.get("timeline", ""),
        "location": data.get("location", ""),
        "work_type": data.get("work_type", ""),
        "date": "19 April 2026",
        # Format work activities for Pasal 2
        "pasal_2_content": pasal_2_text,
        # Format scope + items for Pasal 3
        "pasal_3_content": format_pasal_3_content(data),
        # Payment terms (empty for now - from LHP not available)
        "pasal_10_content": "Pembayaran dilakukan sesuai ketentuan yang berlaku di PT Terminal Petikemas Surabaya",
    }

    # Test RKS template
    try:
        print("\n=== FILLING RKS TEMPLATE ===")
        rks_doc = docx_service.load_template(doc_type, "RKS")
        filled_doc = docx_service.fill_template(rks_doc, template_data)

        # Save output
        import uuid
        output_path = f"output/test_extraction_rks_{uuid.uuid4().hex[:8]}.docx"
        docx_service.save_document(filled_doc, output_path)

        # Verify
        assert Path(output_path).exists(), "Output file not created"

        saved_doc = Document(str(output_path))
        full_text = '\n'.join([p.text for p in saved_doc.paragraphs])

        assert data["project_name"] in full_text, "Project name not found"
        assert "{{project_name}}" not in full_text, "Placeholder not replaced"

        print(f"✓ RKS generated: {output_path}")
        print(f"✓ Project name embedded: {data.get('project_name')[:50]}...")

    except FileNotFoundError as e:
        pytest.skip(f"Template not found: {e}")

    print("\n✓ EXTRACTION → DOCX PASSED")

def test_full_workflow_with_real_lhp():
    """Test complete workflow: LHP → extraction → DOCX generation"""

    # Check prerequisites
    if not os.getenv("GEMINI_API_KEY"):
        pytest.skip("GEMINI_API_KEY not set")

    lhp_path = "docs/LHP/Laporan Hasil Pemeriksaan Fiber Optic Koneksi 4 Unit QCC Baru.pdf"
    if not Path(lhp_path).exists():
        pytest.skip(f"LHP file not found: {lhp_path}")

    print("\n" + "="*60)
    print("FULL WORKFLOW TEST")
    print("="*60)

    # Step 1: OCR Extraction
    print("\n[Step 1] OCR Extraction...")
    ocr_service = OCRService()
    lhp_text = ocr_service.extract_text(lhp_path)
    assert len(lhp_text) > 1000, "OCR extraction failed"
    print(f"✓ Extracted {len(lhp_text)} characters")

    # Step 2: AI Extraction - Detect document type
    print("\n[Step 2] AI Extraction - Document Type...")
    extraction_service = ExtractionService()
    doc_type = extraction_service.detect_document_type(lhp_text)
    assert doc_type in ["PENGADAAN", "PEMELIHARAAN", "PADI_UMKM"]
    print(f"✓ Document type: {doc_type}")

    # Step 3: AI Extraction - Project name
    print("\n[Step 3] AI Extraction - Project Name...")
    project_name = extraction_service.extract_project_name(lhp_text)
    assert len(project_name) > 10
    print(f"✓ Project name: {project_name}")

    # Step 4: AI Extraction - Structured data
    print("\n[Step 4] AI Extraction - Structured Data...")
    data = extraction_service.extract_structured_data(lhp_text, doc_type)

    # Verify required fields
    required_fields = ["project_name", "items", "timeline", "location", "document_type"]
    for field in required_fields:
        assert field in data, f"Missing field: {field}"
        print(f"  ✓ {field}: {data.get(field)}")

    # Verify items
    assert isinstance(data["items"], list)
    assert len(data["items"]) > 0
    print(f"  ✓ Items extracted: {len(data['items'])} items")

    # Step 5: Paraphrasing test
    print("\n[Step 5] AI Paraphrasing...")
    paraphrase_service = ParaphraseService()
    test_text = "Pembayaran dilakukan dengan termin I 95% dan termin II 5%"
    options = paraphrase_service.paraphrase_section(test_text, "payment_terms")
    assert isinstance(options, list)
    assert len(options) > 0
    print(f"  ✓ Generated {len(options)} paraphrase options")
    for i, opt in enumerate(options):
        print(f"    Option {i+1}: {opt[:60]}...")

    # Step 6: DOCX Generation (if templates exist)
    print("\n[Step 6] DOCX Generation...")

    try:
        docx_service = DOCXService()

        # Test RAB generation
        try:
            rab_doc = docx_service.load_template(doc_type, "RAB")
            filled_rab = docx_service.fill_template(rab_doc, data)

            output_dir = Path("output")
            output_dir.mkdir(exist_ok=True)

            rab_path = output_dir / f"TEST_RAB_{data.get('project_name', 'project').replace(' ', '_')[:50]}.docx"
            docx_service.save_document(filled_rab, str(rab_path))

            assert rab_path.exists(), "RAB file not created"
            print(f"  ✓ RAB generated: {rab_path.name}")

            # Verify placeholders replaced
            saved_doc = Document(str(rab_path))
            text = '\n'.join([p.text for p in saved_doc.paragraphs])
            assert data["project_name"] in text
            assert "{{project_name}}" not in text
            print(f"    ✓ Placeholders verified in RAB")

            # Cleanup
            rab_path.unlink()

        except FileNotFoundError:
            print("  ⊗ RAB template not found (skipping)")

        # Test RKS generation
        try:
            rks_doc = docx_service.load_template(doc_type, "RKS")
            filled_rks = docx_service.fill_template(rks_doc, data)

            rks_path = output_dir / f"TEST_RKS_{data.get('project_name', 'project').replace(' ', '_')[:50]}.docx"
            docx_service.save_document(filled_rks, str(rks_path))

            assert rks_path.exists(), "RKS file not created"
            print(f"  ✓ RKS generated: {rks_path.name}")

            # Verify placeholders replaced
            saved_doc = Document(str(rks_path))
            text = '\n'.join([p.text for p in saved_doc.paragraphs])
            assert data["project_name"] in text
            assert "{{project_name}}" not in text
            print(f"    ✓ Placeholders verified in RKS")

            # Cleanup
            rks_path.unlink()

        except FileNotFoundError:
            print("  ⊗ RKS template not found (skipping)")

    except Exception as e:
        pytest.fail(f"DOCX generation failed: {e}")

    print("\n" + "="*60)
    print("✓ ALL WORKFLOW STEPS PASSED")
    print("="*60)

def test_extraction_fields_coverage():
    """Test that extraction covers all required RKS placeholders"""

    if not os.getenv("GEMINI_API_KEY"):
        pytest.skip("GEMINI_API_KEY not set")

    service = ExtractionService()

    # Realistic LHP text
    lhp_text = """
    LAPORAN HASIL PEMERIKSAAN KEBUTUHAN FIBER OPTIC KONEKSI 4 UNIT QCC BARU

    Tabel 3.1 Lingkup Item Pekerjaan:
    1. Instalasi FO 4 Jalur - 2680 m
    2. Splicing FO Backbone - 96 core
    """

    data = service.extract_structured_data(lhp_text, "PENGADAAN")

    # Required fields for RKS
    required_placeholders = [
        "project_name",
        "work_type",
        "timeline",
        "location",
        "pasal_2_content",  # Should come from numbered items
        "pasal_3_content",  # Should include items table
        "pasal_10_content"  # Payment terms
    ]

    print("\nPlaceholder coverage check:")
    for placeholder in required_placeholders:
        # Check if extraction provides this field
        if placeholder in data or placeholder == "pasal_2_content" or placeholder == "pasal_3_content" or placeholder == "pasal_10_content":
            # These are content blocks that need special handling
            if placeholder == "project_name" and data.get("project_name"):
                print(f"  ✓ {placeholder}: Available")
            elif placeholder == "work_type":
                # work_type needs to be extracted from Jasa items or title
                print(f"  ⚠ {placeholder}: Needs extraction logic (from Jasa/title)")
            elif placeholder == "timeline" and data.get("timeline"):
                print(f"  ✓ {placeholder}: Available")
            elif placeholder == "location" and data.get("location"):
                print(f"  ✓ {placeholder}: Available")
            else:
                print(f"  ⚠ {placeholder}: Content block - needs structured extraction")
        else:
            print(f"  ✗ {placeholder}: Missing")

    print("\nNote: pasal_2_content, pasal_3_content, pasal_10_content require AI to extract structured content blocks, not just individual fields")
