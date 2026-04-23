"""Unit tests for insert_numbered_list functionality"""

import pytest
from docx import Document
from services.docx_service import DOCXService


def test_insert_numbered_list_empty_activities(tmp_path):
    """Test that empty activities list is handled gracefully"""
    doc = Document()
    p = doc.add_paragraph("Before {{work_activities}} after")

    service = DOCXService()
    service.insert_numbered_list(doc, [], "{{work_activities}}")

    # Placeholder should still exist but empty
    assert "{{work_activities}}" in p.text or p.text == ""


def test_insert_numbered_list_missing_placeholder(tmp_path):
    """Test behavior when placeholder doesn't exist"""
    doc = Document()
    doc.add_paragraph("Some paragraph without placeholder")

    service = DOCXService()
    service.insert_numbered_list(doc, ["Activity 1"], "{{work_activities}}")

    # Should not crash, just log warning
    assert len(doc.paragraphs) == 1  # No new paragraphs added


def test_insert_numbered_list_with_activities(tmp_path):
    """Test successful insertion of numbered list"""
    doc = Document()
    doc.add_paragraph("Before {{work_activities}}")

    service = DOCXService()
    activities = ["Activity 1", "Activity 2", "Activity 3"]
    service.insert_numbered_list(doc, activities, "{{work_activities}}")

    # Find numbered paragraphs (should have 3 activities)
    text_content = [p.text for p in doc.paragraphs]
    assert "Activity 1" in text_content
    assert "Activity 2" in text_content
    assert "Activity 3" in text_content

    # Check that placeholder is removed
    assert "{{work_activities}}" not in " ".join(text_content)


def test_insert_numbered_list_with_empty_activity(tmp_path):
    """Test that empty strings in activities are skipped"""
    doc = Document()
    doc.add_paragraph("Before {{work_activities}}")

    service = DOCXService()
    activities = ["Activity 1", "", "Activity 2", "   ", "Activity 3"]
    service.insert_numbered_list(doc, activities, "{{work_activities}}")

    # Should only have 3 non-empty activities
    text_content = [p.text for p in doc.paragraphs]
    assert "Activity 1" in text_content
    assert "Activity 2" in text_content
    assert "Activity 3" in text_content

    # Count non-empty activity paragraphs
    activity_paragraphs = [p for p in doc.paragraphs if "Activity" in p.text]
    assert len(activity_paragraphs) == 3
