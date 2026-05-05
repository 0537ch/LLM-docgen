"""
Test to understand python-docx table merging behavior
Focus on vertical merge (vMerge) for DOCX summary table
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.table import Table
from docx.oxml import parse_xml


def test_horizontal_merge():
    """Test horizontal merge (current implementation)"""
    doc = Document()
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'

    # Set headers
    headers = ['NO', 'URAIAN', 'VOLUME', 'SATUAN', 'HARGA SATUAN', 'JUMLAH HARGA']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header

    # Add data rows
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Item A'
    table.rows[1].cells[2].text = '2'
    table.rows[1].cells[3].text = 'unit'
    table.rows[1].cells[4].text = '10000'
    table.rows[1].cells[5].text = '20000'

    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = 'Item B'
    table.rows[2].cells[2].text = '3'
    table.rows[2].cells[3].text = 'kg'
    table.rows[2].cells[4].text = '5000'
    table.rows[2].cells[5].text = '15000'

    # Merge cells 0-4 in row 3 (summary row - horizontal merge)
    row = table.rows[3]
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = "Total"
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row.cells[5].text = "35000"

    # Row 4
    row4 = table.rows[4]
    row4.cells[0].merge(row4.cells[4])
    row4.cells[0].text = "PPN (11%)"
    row4.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row4.cells[5].text = "3850"

    output_path = "output_test/test_horizontal_merge.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


def test_vertical_merge():
    """Test vertical merge approach"""
    doc = Document()
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'

    # Set headers
    headers = ['NO', 'URAIAN', 'VOLUME', 'SATUAN', 'HARGA SATUAN', 'JUMLAH HARGA']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header

    # Add data rows
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Item A'
    table.rows[1].cells[2].text = '2'
    table.rows[1].cells[3].text = 'unit'
    table.rows[1].cells[4].text = '10000'
    table.rows[1].cells[5].text = '20000'

    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = 'Item B'
    table.rows[2].cells[2].text = '3'
    table.rows[2].cells[3].text = 'kg'
    table.rows[2].cells[4].text = '5000'
    table.rows[2].cells[5].text = '15000'

    # Attempt: merge horizontally in each row (col 0-2), then add label in col 3
    for row_idx in [3, 4]:
        row = table.rows[row_idx]
        row.cells[0].merge(row.cells[1])
        row.cells[0].merge(row.cells[2])

    table.rows[3].cells[0].text = "Terbilang: satu miliar rupiah"
    table.rows[3].cells[3].text = "Total"
    table.rows[3].cells[5].text = "35000"

    table.rows[4].cells[0].text = ""
    table.rows[4].cells[3].text = "PPN (11%)"
    table.rows[4].cells[5].text = "3850"

    output_path = "output_test/test_vertical_merge.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


def test_vmerge_same_column():
    """Try to understand vMerge in same column across rows"""
    doc = Document()
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'

    # Set simple content
    table.rows[0].cells[0].text = "A"
    table.rows[1].cells[0].text = "B"
    table.rows[2].cells[0].text = "C"
    table.rows[3].cells[0].text = "D"

    # Inspect the tc (table cell) XML for vMerge attributes
    for row_idx, row in enumerate(table.rows):
        tc = row.cells[0]._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vMerge = tcPr.find(qn('w:vMerge'))
            print(f"Row {row_idx} vMerge: {vMerge}")
        else:
            print(f"Row {row_idx}: no tcPr found")

    output_path = "output_test/test_vmerge_basic.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


def set_vmerge(cell, val):
    """Set vMerge attribute on cell. val='restart' starts merge, 'continue' continues, None removes"""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    vMerge = tcPr.find(qn('w:vMerge'))
    if vMerge is None:
        vMerge = OxmlElement('w:vMerge')
        tcPr.append(vMerge)
    if val is None:
        tcPr.remove(vMerge)
    else:
        vMerge.set(qn('w:val'), val)


def test_combined_vertical_horizontal():
    """Test: horizontal merge in rows, then vertical merge across those merged cells using XML"""
    doc = Document()
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'

    # Headers
    headers = ['NO', 'URAIAN', 'VOLUME', 'SATUAN', 'HARGA SATUAN', 'JUMLAH HARGA']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data
    table.rows[1].cells[0].text = '1'
    table.rows[1].cells[1].text = 'Item A'
    table.rows[1].cells[2].text = '2'
    table.rows[1].cells[3].text = 'unit'
    table.rows[1].cells[4].text = '10000'
    table.rows[1].cells[5].text = '20000'

    table.rows[2].cells[0].text = '2'
    table.rows[2].cells[1].text = 'Item B'
    table.rows[2].cells[2].text = '3'
    table.rows[2].cells[3].text = 'kg'
    table.rows[2].cells[4].text = '5000'
    table.rows[2].cells[5].text = '15000'

    # Row 3: Total row
    table.rows[3].cells[0].merge(table.rows[3].cells[1])
    table.rows[3].cells[0].merge(table.rows[3].cells[2])
    table.rows[3].cells[0].text = ""  # will be vertically merged
    table.rows[3].cells[3].text = "Total"
    table.rows[3].cells[5].text = "35000"

    # Row 4: PPN row
    table.rows[4].cells[0].merge(table.rows[4].cells[1])
    table.rows[4].cells[0].merge(table.rows[4].cells[2])
    table.rows[4].cells[0].text = ""  # will be vertically merged
    table.rows[4].cells[3].text = "PPN (11%)"
    table.rows[4].cells[5].text = "3850"

    # Now apply vMerge to make the merged cells span vertically
    # Row 3 col 0 = restart, Row 4 col 0 = continue
    set_vmerge(table.rows[3].cells[0], 'restart')
    set_vmerge(table.rows[4].cells[0], 'continue')

    output_path = "output_test/test_combined.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    import os
    os.makedirs("output_test", exist_ok=True)

    print("=== Test 1: Horizontal Merge (current implementation) ===")
    test_horizontal_merge()

    print("\n=== Test 2: Vertical Merge attempt ===")
    test_vertical_merge()

    print("\n=== Test 3: vMerge same column inspection ===")
    test_vmerge_same_column()

    print("\n=== Test 4: Combined vertical + horizontal ===")
    test_combined_vertical_horizontal()