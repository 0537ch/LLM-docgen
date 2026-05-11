"""
Test: Parse RKS DOCX template with python-docx directly, extract content + alignment, build HTML preview.
Goal: Replace mammoth conversion with manual HTML generation from parsed DOCX.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Paths
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_TEST_DIR = Path(__file__).parent.parent / "output_test"
OUTPUT_TEST_DIR.mkdir(exist_ok=True)

def parse_docx_to_html(template_path: Path) -> str:
    """
    Parse DOCX with python-docx, extract text + alignment, build clean HTML.
    """
    doc = Document(str(template_path))

    html_parts = []
    html_parts.append("""
    <style>
        body { font-family: 'Times New Roman', serif; font-size: 12pt; margin: 20px; }
        p { margin: 0; text-align: justify; line-height: 1.5; }
        .center { text-align: center; }
        .right { text-align: right; }
        table { border-collapse: collapse; width: 100%; margin: 16px 0; }
        td, th { border: 1px solid #333; padding: 6px 8px; }
        th { background-color: #f0f0f0; text-align: center; }
        td { text-align: left; }
        h1, h2, h3 { text-align: center; margin: 16px 0 8px 0; }
        ol, ul { margin: 0; padding-left: 24px; }
        .item-name { text-align: left; }
        .item-qty { text-align: center; }
    </style>
    """)

    html_parts.append("<body>")

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]  # Strip namespace

        if tag == 'p':
            # Get paragraph
            para = element
            p_text = ""
            p_alignment = None

            # Extract text and alignment from paragraph
            for child in element.iter():
                child_tag = child.tag.split('}')[-1] if child.tag else ''
                if child_tag == 'r':
                    # Text run
                    for t in child.iter():
                        if t.tag.split('}')[-1] == 't' and t.text:
                            p_text += t.text
                elif child_tag == 'pPr':
                    # Paragraph properties
                    for jc in child.iter():
                        if jc.tag.split('}')[-1] == 'jc':
                            val = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            p_alignment = val

            # Determine alignment class
            align_class = ""
            if p_alignment == 'center':
                align_class = "center"
            elif p_alignment == 'right':
                align_class = "right"

            # Build HTML paragraph
            if p_text.strip():
                if align_class:
                    html_parts.append(f'<p class="{align_class}">{p_text}</p>')
                else:
                    html_parts.append(f'<p>{p_text}</p>')

        elif tag == 'tbl':
            # Handle table
            tbl = Document(None).element.body  # placeholder
            html_parts.append("<!-- TABLE -->")

        elif tag == 'sectPr':
            pass  # Section properties, skip

    html_parts.append("</body>")
    return "\n".join(html_parts)


def test_rks_parsing():
    """Test parsing RKS pengadaan template"""
    template_path = TEMPLATES_DIR / "RKS_pengadaan.docx"

    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return

    doc = Document(str(template_path))

    print(f"\n=== Parsing {template_path} ===\n")

    # Check all paragraph alignments
    print("Paragraphs with alignment info:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        alignment = para.alignment

        align_name = "left"
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_name = "CENTER"
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_name = "RIGHT"
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            align_name = "JUSTIFY"

        print(f"  [{i:3d}] [{align_name:8s}] {text[:80]}")

    # Check tables
    print(f"\nTables found: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"  Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")


    # Generate HTML preview
    html = parse_docx_to_html(template_path)
    output_path = OUTPUT_TEST_DIR / "rks_direct_parse.html"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

    print(f"\nHTML saved: {output_path}")


if __name__ == '__main__':
    test_rks_parsing()
