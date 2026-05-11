from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
from typing import Dict, Any, List
from utils.config import Config
from utils.logger import setup_logger

logger = setup_logger("docx_service")

class DOCXService:
    def __init__(self):
        self.templates_dir = Path(Config.TEMPLATES_DIR)

    def load_template(self, doc_type: str, template_type: str) -> Document:
        """Load DOCX template"""
        template_path = self.templates_dir / f"{template_type}_{doc_type.lower()}.docx"

        if not template_path.exists():
            # Try alternative naming
            template_path = self.templates_dir / f"{doc_type}_{template_type}.docx"

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(str(template_path))
        return doc

    def add_items_table(self, doc: Document, items: List[Dict], table_title: str = '', placeholder: str = None) -> 'docx.table.Table':
        """Add items table with actual DOCX table structure

        Args:
            doc: Document object
            items: List of item dictionaries
            table_title: Optional title for the table
            placeholder: If provided, find this placeholder and insert table there instead of at end

        Returns:
            The created table object
        """
        if not items:
            logger.warning("No items to add to table")
            return

        # Find insertion point if placeholder specified
        placeholder_paragraph = None

        if placeholder:
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    placeholder_paragraph = paragraph
                    logger.info(f"Found placeholder '{placeholder}'")
                    break

        # Create table using python-docx API (6 columns: NO, URAIAN, VOLUME, SATUAN, HARGA SATUAN, JUMLAH HARGA)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Set table width to auto for autofit to CONTENT (not page width)
        from docx.oxml.shared import OxmlElement, qn
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:type'), 'auto')
        tbl_w.set(qn('w:w'), '0')
        table._element.tblPr.append(tbl_w)

        # Enable autofit for table
        table.autofit = True

        # Remove fixed column widths to allow autofit to work
        for column in table.columns:
            column.width = None

        # Add headers
        headers = ['NO', 'URAIAN', 'VOLUME', 'SATUAN', 'HARGA SATUAN', 'JUMLAH HARGA']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.bold = False

        # Add data rows
        for idx, item in enumerate(items, start=1):
            row_cells = table.add_row().cells

            # NO
            row_cells[0].text = str(item.get('NO', idx))
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[0].paragraphs[0].paragraph_format.line_spacing = 1.0

            # URAIAN
            uraian = item.get('uraian', item.get('name', item.get('Deskripsi', item.get('URAIAN', ''))))
            row_cells[1].text = str(uraian)
            for paragraph in row_cells[1].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0

            # VOLUME
            volume = item.get('volume', item.get('quantity', item.get('Volume', item.get('VOLUME', ''))))
            row_cells[2].text = str(volume)
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[2].paragraphs[0].paragraph_format.line_spacing = 1.0

            # SATUAN
            satuan = item.get('satuan', item.get('unit', item.get('Satuan', item.get('SATUAN', ''))))
            row_cells[3].text = str(satuan)
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].paragraph_format.line_spacing = 1.0

            # HARGA SATUAN
            harga = item.get('harga_satuan', item.get('price', item.get('Harga', item.get('HARGA', ''))))
            try:
                harga_val = float(harga) if harga else 0
                harga_text = f"Rp {int(harga_val):,}".replace(',', '.') if harga_val > 0 else '-'
            except (ValueError, TypeError):
                harga_text = str(harga) if harga else '-'
            row_cells[4].text = harga_text
            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[4].paragraphs[0].paragraph_format.line_spacing = 1.0

            # JUMLAH HARGA (volume * harga_satuan)
            try:
                vol = float(item.get('volume', 0)) if item.get('volume') else 0
                sat = float(item.get('harga_satuan', 0)) if item.get('harga_satuan') else 0
                jumlah = vol * sat
                jumlah_text = f"Rp {int(jumlah):,}".replace(',', '.') if jumlah > 0 else '-'
            except (ValueError, TypeError):
                jumlah_text = '-'
            row_cells[5].text = jumlah_text
            row_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[5].paragraphs[0].paragraph_format.line_spacing = 1.0

        # Move table to placeholder position if specified
        if placeholder_paragraph:
            # Get table element
            table_element = table._element

            # Insert table element after placeholder paragraph
            placeholder_paragraph._element.addnext(table_element)

            # Clear placeholder text but keep paragraph
            placeholder_paragraph.text = ""
            logger.info(f"Moved table to placeholder position")
        else:
            logger.info("Appended table to end of document")

        return table

    def add_items_table_no_price(self, doc: Document, items: List[Dict], table_title: str = '', placeholder: str = None) -> 'docx.table.Table':
        """Add items table WITHOUT price columns (for RKS Pasal 3)

        Args:
            doc: Document object
            items: List of item dictionaries
            table_title: Optional title for the table
            placeholder: If provided, find this placeholder and insert table there instead of at end

        Returns:
            The created table object
        """
        if not items:
            logger.warning("No items to add to table")
            return

        # Find insertion point if placeholder specified
        placeholder_paragraph = None

        if placeholder:
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    placeholder_paragraph = paragraph
                    logger.info(f"Found placeholder '{placeholder}'")
                    break

        # Create table using python-docx API (4 columns, no price)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Enable autofit for table
        table.autofit = True

        # Remove fixed column widths to allow autofit to work
        for column in table.columns:
            column.width = None

        # Add headers (no HARGA SATUAN)
        headers = ['NO', 'URAIAN', 'VOLUME', 'SATUAN']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.bold = False

        # Add data rows (only 4 columns)
        for idx, item in enumerate(items, start=1):
            row_cells = table.add_row().cells

            # NO
            row_cells[0].text = str(item.get('NO', idx))
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[0].paragraphs[0].paragraph_format.line_spacing = 1.0

            # URAIAN
            uraian = item.get('uraian', item.get('name', item.get('Deskripsi', item.get('URAIAN', ''))))
            row_cells[1].text = str(uraian)
            for paragraph in row_cells[1].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0

            # VOLUME
            volume = item.get('volume', item.get('quantity', item.get('Volume', item.get('VOLUME', ''))))
            row_cells[2].text = str(volume)
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[2].paragraphs[0].paragraph_format.line_spacing = 1.0

            # SATUAN
            satuan = item.get('satuan', item.get('unit', item.get('Satuan', item.get('SATUAN', ''))))
            row_cells[3].text = str(satuan)
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].paragraph_format.line_spacing = 1.0

        # Move table to placeholder position if specified
        if placeholder_paragraph:
            # Get table element
            table_element = table._element

            # Insert table element after placeholder paragraph
            placeholder_paragraph._element.addnext(table_element)

            # Clear placeholder text but keep paragraph
            placeholder_paragraph.text = ""
            logger.info(f"Moved table to placeholder position")
        else:
            logger.info("Appended table to end of document")
        logger.info(f"Added table with {len(items)} rows")
        return table

    def add_summary_table(self, doc: Document, table: 'docx.table.Table', ppn_percent: int = 11) -> None:
        """Add summary rows (Total, PPN, Grand Total) to the items table

        Args:
            doc: Document object
            table: The items table to add summary to
            ppn_percent: PPN percentage (default 11%)
        """
        # Calculate total from existing rows
        total = 0
        for row in table.rows[1:]:  # Skip header row
            try:
                # Jumlah Harga is in column 5 (0-indexed), after NO|URAIAN|VOLUME|SATUAN|HARGA SATUAN|JUMLAH HARGA
                jumlah_text = row.cells[5].text.strip()
                if jumlah_text and jumlah_text != '-':
                    # Remove non-numeric characters (dots, commas, spaces, 'Rp', etc)
                    jumlah_clean = ''.join(c for c in jumlah_text if c.isdigit())
                    if jumlah_clean:
                        total += int(jumlah_clean)
            except (ValueError, IndexError):
                continue

        # Calculate PPN and Grand Total
        ppn = int(total * ppn_percent / 100)
        grand_total = total + ppn

        # Add Total row (merge cells 0-4, place total in cell 5)
        total_row = table.add_row()
        total_row.cells[0].merge(total_row.cells[4])  # Merge first 5 cells (NO through HARGA SATUAN)
        total_row.cells[0].text = "Total"
        total_row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        total_row.cells[0].paragraphs[0].runs[0].bold = True
        total_row.cells[5].text = f"Rp {total:,}".replace(',', '.')
        total_row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add PPN row
        ppn_row = table.add_row()
        ppn_row.cells[0].merge(ppn_row.cells[4])
        ppn_row.cells[0].text = f"PPN ({ppn_percent}%)"
        ppn_row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        ppn_row.cells[0].paragraphs[0].runs[0].bold = True
        ppn_row.cells[5].text = f"Rp {ppn:,}".replace(',', '.')
        ppn_row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add Grand Total row
        grand_row = table.add_row()
        grand_row.cells[0].merge(grand_row.cells[4])
        grand_row.cells[0].text = "Grand Total"
        grand_row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        grand_row.cells[0].paragraphs[0].runs[0].bold = True
        grand_row.cells[5].text = f"Rp {grand_total:,}".replace(',', '.')
        grand_row.cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        logger.info(f"Added summary: Total={total}, PPN={ppn}, Grand Total={grand_total}")

    def replace_placeholders(self, doc: Document, replacements: Dict[str, Any], list_placeholders: List[str] = None) -> None:
        """Replace placeholders in document while preserving formatting

        Args:
            doc: Document object
            replacements: Dict of placeholder names to values (strings or lists)
            list_placeholders: List of placeholder keys that should be formatted as numbered lists
        """
        if list_placeholders is None:
            list_placeholders = []

        logger.info(f"Replacing placeholders: {list(replacements.keys())}")
        logger.info(f"List placeholders: {list_placeholders}")
        replaced_count = 0

        # Handle list-type placeholders first
        for list_key in list_placeholders:
            if list_key in replacements:
                placeholder = f"{{{{{list_key}}}}}"
                value = replacements[list_key]

                if isinstance(value, list):
                    self.insert_numbered_list(doc, value, placeholder)
                    replaced_count += 1
                else:
                    logger.warning(f"List placeholder '{list_key}' received non-list value: {type(value)}")

        # Handle normal text placeholders (skip list placeholders)
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in list_placeholders:
                    continue  # Already handled

                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, placeholder, str(value))
                    logger.info(f"Replaced placeholder: {placeholder}")
                    replaced_count += 1

        # Also check tables for normal placeholders
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in list_placeholders:
                                continue  # Already handled

                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, placeholder, str(value))
                                logger.info(f"Replaced placeholder in table: {placeholder}")
                                replaced_count += 1

        logger.info(f"Total placeholders replaced: {replaced_count}")

    def _replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """Replace text in paragraph while preserving formatting"""
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    def fill_template(self, doc: Document, data: Dict[str, Any], list_placeholders: List[str] = None) -> Document:
        """Fill template with extracted data

        Args:
            doc: Document object
            data: Dict of placeholder names to values
            list_placeholders: List of keys that should be formatted as numbered lists
        """
        if list_placeholders is None:
            list_placeholders = []

        # Convert all values to strings except list placeholders
        replacements = {}
        for key, value in data.items():
            if key in list_placeholders:
                # Keep as-is (should be list)
                replacements[key] = value
            else:
                # Convert to string, None -> empty string
                replacements[key] = str(value) if value is not None else ""

        self.replace_placeholders(doc, replacements, list_placeholders)
        return doc


    def save_document(self, doc: Document, output_path: str) -> None:
        """Save document to file"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Saving document to: {output_path}")
        doc.save(str(output_path))

    def docx_to_html(self, doc: Document) -> str:
        """Convert docx Document to HTML string using mammoth"""
        from mammoth import convert_to_html
        from bs4 import BeautifulSoup

        # Save docx to bytes for mammoth
        from io import BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Convert to HTML
        result = convert_to_html(docx_bytes)

        # Post-process: center paragraphs that contain <strong> tags
        soup = BeautifulSoup(result.value, 'html.parser')
        for p in soup.find_all('p'):
            if p.find('strong'):
                p['style'] = 'text-align: center; font-weight: bold; margin: 12px 0'
                # Keep the strong tag inside, just center the paragraph
            # Right-align "Dikeluarkan di" and "Tanggal" paragraphs
            text = p.get_text().strip()
            if text.startswith("Dikeluarkan di") or text.startswith("Tanggal:"):
                p['style'] = 'text-align: right'

        # Wrap with CSS for proper table styling
        styled_html = f"""
        <style>
            .rks-preview {{ font-family: 'Times New Roman', serif; font-size: 12pt; margin: 0 1cm; padding: 0; }}
            .rks-preview table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
            .rks-preview td, .rks-preview th {{ border: 1px solid #333; padding: 8px; vertical-align: top; }}
            .rks-preview th {{ background-color: #f0f0f0; text-align: center; }}
            .rks-preview td {{ text-align: left; }}
            .rks-preview h1, .rks-preview h2, .rks-preview h3 {{ text-align: center; margin: 0; }}
            .rks-preview p {{ margin: 0; text-align: justify; }}
            .rks-preview p[style*='text-align: center'] {{ text-align: center !important; }}
            .rks-preview .item-name {{ text-align: left !important; }}
            .rks-preview .item-qty {{ text-align: center !important; }}
            .rks-preview ol, .rks-preview ul {{ margin: 0; padding-left: 24px; }}
        </style>
        <div class="rks-preview">
        {str(soup)}
        </div>
        """
        return styled_html

    def insert_numbered_list(self, doc: Document, activities: List[str], placeholder: str) -> None:
        """Insert work activities as Word numbered list at placeholder position

        Args:
            doc: Document object
            activities: List of activity strings (no manual numbering)
            placeholder: Placeholder text to find (e.g., "{{pasal2_content}}")

        Behavior:
            - Finds paragraph containing placeholder
            - Inserts each activity as separate paragraph with "List Number" style
            - Deletes placeholder paragraph
            - Falls back to manual numbering if "List Number" style missing
        """
        if not activities:
            logger.warning("No activities to insert")
            return

        # Find placeholder paragraph
        placeholder_paragraph = None
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_paragraph = paragraph
                logger.info(f"Found placeholder '{placeholder}'")
                break

        if not placeholder_paragraph:
            logger.warning(f"Placeholder '{placeholder}' not found")
            return

        # Check if "List Number" style exists
        style_name = "List Number"
        style_exists = any(s.name == style_name for s in doc.styles)

        # Get placeholder position (insertion point)
        placeholder_element = placeholder_paragraph._element

        # Store placeholder formatting for copying
        placeholder_alignment = placeholder_paragraph.alignment
        placeholder_paragraph_format = placeholder_paragraph.paragraph_format

        # Get font from placeholder (use first run if available)
        placeholder_font = None
        if placeholder_paragraph.runs:
            placeholder_font = placeholder_paragraph.runs[0].font

        # Track insertion point to maintain order
        last_element = placeholder_element

        # Insert each activity as numbered paragraph
        for idx, activity in enumerate(activities, start=1):
            if not activity.strip():
                continue

            # Create new paragraph
            new_para = doc.add_paragraph(activity)

            # Copy placeholder formatting
            new_para.alignment = placeholder_alignment

            # Copy paragraph format (line spacing, etc.)
            if placeholder_paragraph_format:
                new_para.paragraph_format.line_spacing = placeholder_paragraph_format.line_spacing
                new_para.paragraph_format.space_before = placeholder_paragraph_format.space_before
                new_para.paragraph_format.space_after = placeholder_paragraph_format.space_after

            # Copy font formatting
            if placeholder_font and new_para.runs:
                new_font = new_para.runs[0].font
                new_font.name = placeholder_font.name
                new_font.size = placeholder_font.size
                new_font.bold = placeholder_font.bold
                new_font.italic = placeholder_font.italic

            # Apply numbering style or fallback
            if style_exists:
                try:
                    new_para.style = style_name
                except Exception as e:
                    logger.warning(f"Failed to apply style: {e}, using manual numbering")
                    # Add manual numbering as fallback
                    new_para.text = f"{idx}. {activity}"
            else:
                # Manual numbering fallback
                new_para.text = f"{idx}. {activity}"
                logger.warning(f"Style '{style_name}' not found, using manual numbering")

            # Move paragraph to correct position (after last inserted)
            new_para_element = new_para._element
            last_element.addnext(new_para_element)
            last_element = new_para_element

        # Remove placeholder paragraph
        placeholder_paragraph.text = ""
        logger.info(f"Inserted {len(activities)} activities as numbered list")

